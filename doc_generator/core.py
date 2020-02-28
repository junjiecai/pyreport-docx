#!/usr/bin/env python
# coding: utf-8

# # todo
# * 开始处理模版模式
# * 补全方法
# * 完善类的设计
# * 支持自定义组件

# In[1]:
from abc import ABC
from typing import List as ListType

from docx import Document
import re


def _attach_table(document, headers, data):
    table = document.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells

    for i, text in enumerate(headers):
        hdr_cells[i].text = text

    for record in data:
        row_cells = table.add_row().cells

        for i, value in enumerate(record):
            row_cells[i].text = str(value)


class Component:
    def __init__(self):
        self.components: ListType[Component] = []

    def _attach(self, parent_component):
        raise NotImplementedError


class Paragraph(Component):
    def __init__(self, text):
        self.text = text

    def _attach(self, parent_component):
        parent_component.add_paragraph(self.text)


class Header(Component):
    def __init__(self, text, level=1):
        self.text = text
        self.level = level

    def _attach(self, parent_component):
        parent_component.add_heading(self.text, level=self.level)


class Image(Component):
    def __init__(self, path):
        self.path = path

    def _attach(self, parent_component):
        parent_component.add_picture(self.path)


class Table(Component):
    def __init__(self, headers, data):
        self.headers = headers
        self.data = data

    def _attach(self, parent_component):
        _attach_table(parent_component, self.headers, self.data)


class List(Component, ABC):
    def _add_item(self, item):
        self.components.append(item)


class ListItem(Component):
    def __init__(self, text, parent_component):
        Component.__init__(self)
        self.text = text
        self.parent_component = parent_component

    def _attach(self, parent_component):
        level = self.parent_component.level
        if level > 1:
            parent_component.add_paragraph(self.text, style="List Number {}".format(level))
        else:
            parent_component.add_paragraph(self.text, style="List Number")


class OrderedList(List):
    def _add_item(self, item):
        pass

    def __init__(self, texts, level=1):
        List.__init__(self)
        self.texts = texts
        self.level = level
        self.initialize_components()

    def initialize_components(self):
        for text in self.texts:
            if isinstance(text, str):
                self.components.append(ListItem(text, self))
            else:
                self.transmit_level(text)
                self.components.append(text)

    def transmit_level(self, text):
        text.level = self.level + 1
        for component in text.components:
            if isinstance(component, OrderedList):
                component.transmit_level(component)

    def _attach(self, parent_component):
        for component in self.components:
            component._attach(parent_component)


class Doc:
    def __init__(self):
        self.components: ListType[Component] = []

    def add_header(self, text, level=1):
        self.components.append(Header(text, level=level))

    def add_paragraph(self, text):
        self.components.append(Paragraph(text))

    def add_ordered_list(self, texts):
        self.components.append(OrderedList(texts))

    def add_unordered_list(self):
        pass

    def add_image(self, path):
        self.components.append(Image(path))

    def add_table(self, headers, data):
        self.components.append(Table(headers, data))

    def to_html(self):
        pass

    def to_docx(self, file_name):
        docx = Document()

        for component in self.components:
            component._attach(docx)

        docx.save(file_name)


def _get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment


def _delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def _get_tag(text):
    result = re.search("{{(.*)}}", text)

    if result:
        tag = result.groups()[0]

        tag_type, var_name = map(lambda s: s.strip(), tag.split(":"))

        return tag_type, var_name
    else:
        return None, None


def render(document, document_des, header_data, image_data, list_data, save_path, table_data, text_data):
    for p in document_des.paragraphs:
        _delete_paragraph(p)
    for p in document.paragraphs:
        text = p.text

        tag_type, tag_name = _get_tag(text)

        if tag_type:
            if tag_type == "text":
                document_des.add_paragraph(text_data.get(tag_name))
            elif tag_type == "img":
                document_des.add_picture(image_data.get(tag_name))
            elif tag_type.startswith("header"):
                level = int(tag_type.split("_")[1])
                document_des.add_heading(header_data.get(tag_name), level=level)
            elif tag_type == "table":
                table = Table(table_data.get(tag_name)['headers'], table_data.get(tag_name)['data'])
                table._attach(document_des)
            elif tag_name == 'list':
                l = OrderedList(list_data.get(tag_name))
                l._attach(document_des)
            else:
                _get_para_data(document_des, p)
        else:
            _get_para_data(document_des, p)
    document_des.save(save_path)